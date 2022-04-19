'''
Author: Wei-Ju-Chen(Idonotwant)
ProjectName: py_word
Description: py_word
Brief_Log:
2022/04/19 Start
Note: 
'''
from os import system
from docx import Document
from docx.shared import Inches

#create new empty document/add author
document = Document()
document.core_properties.author = "Chen-Wei-Ju"

#add heading(just type 0 is beautifulQQ)
i = 0
#for i in range(10):
document.add_heading(f'DocumentTitle{i}',i)



#add paragraph
p1 = document.add_paragraph(text="paragraph1")
#add run in p1
p1.add_run("run0").add_tab()
p1.add_run("run1,with italic").italic = True
p1.add_run("run1.5").add_break()#support many types,default is for line
p1.add_run("run2,after break").underline = True
p1.add_run("run3,").Font = 'Microsoft YaHei'
document.add_paragraph(text="paragraph2")

#add page break
document.add_page_break()

#save
document.save('test1.docx')

system(' test1.docx')